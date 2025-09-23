"""
Educational Q&A System with RAG (Retrieval-Augmented Generation)

This cog provides:
- /ask command for technical questions
- Document indexing from /content directory
- LlamaIndex integration for retrieval
- ChromaDB vector storage
- Conversation memory for follow-ups
- Source citations
- Rate limiting integration
"""

import os
import asyncio
import logging
from typing import Dict, List, Optional, Any
from pathlib import Path
import json
from datetime import datetime, timedelta

import discord
from discord.ext import commands
import aiosqlite

# RAG and Document Processing
from llama_index.core import (
    VectorStoreIndex, 
    SimpleDirectoryReader, 
    StorageContext,
    Settings,
    ChatPromptTemplate
)
from llama_index.vector_stores.chroma import ChromaVectorStore
from llama_index.llms.openai import OpenAI
from llama_index.embeddings.openai import OpenAIEmbedding
import chromadb
from chromadb.config import Settings as ChromaSettings

# Document processing
import pypdf
from docx import Document as DocxDocument
import pdfplumber

logger = logging.getLogger(__name__)

class ConversationMemory:
    """Manages conversation context for follow-up questions"""
    
    def __init__(self, max_history: int = 5):
        self.conversations: Dict[int, List[Dict[str, Any]]] = {}
        self.max_history = max_history
    
    def add_interaction(self, user_id: int, question: str, answer: str, sources: List[str]):
        """Add a Q&A interaction to user's conversation history"""
        if user_id not in self.conversations:
            self.conversations[user_id] = []
        
        interaction = {
            "timestamp": datetime.utcnow().isoformat(),
            "question": question,
            "answer": answer,
            "sources": sources
        }
        
        self.conversations[user_id].append(interaction)
        
        # Keep only the most recent interactions
        if len(self.conversations[user_id]) > self.max_history:
            self.conversations[user_id] = self.conversations[user_id][-self.max_history:]
    
    def get_context(self, user_id: int) -> str:
        """Get conversation context for a user"""
        if user_id not in self.conversations:
            return ""
        
        context_parts = []
        for interaction in self.conversations[user_id][-3:]:  # Last 3 interactions
            context_parts.append(f"Previous Q: {interaction['question']}")
            context_parts.append(f"Previous A: {interaction['answer'][:200]}...")
        
        return "\n".join(context_parts)
    
    def clear_conversation(self, user_id: int):
        """Clear conversation history for a user"""
        if user_id in self.conversations:
            del self.conversations[user_id]

class DocumentProcessor:
    """Handles processing of various document formats"""
    
    @staticmethod
    def process_pdf(file_path: Path) -> str:
        """Extract text from PDF files"""
        try:
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                return text
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            return ""
    
    @staticmethod
    def process_docx(file_path: Path) -> str:
        """Extract text from DOCX files"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            return ""
    
    @staticmethod
    def process_text_file(file_path: Path) -> str:
        """Process plain text and markdown files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            return ""

class QACog(commands.Cog):
    """Educational Q&A system with RAG capabilities"""
    
    def __init__(self, bot):
        self.bot = bot
        self.conversation_memory = ConversationMemory()
        self.index = None
        self.query_engine = None
        self.chroma_client = None
        self.is_initializing = False
        
        # Initialize the system
        asyncio.create_task(self.initialize_rag_system())
    
    async def initialize_rag_system(self):
        """Initialize the RAG system with document indexing"""
        if self.is_initializing:
            return
        
        self.is_initializing = True
        
        try:
            logger.info("Initializing RAG system...")
            
            # Configure LlamaIndex settings
            Settings.llm = OpenAI(
                model="gpt-4",
                api_key=os.getenv("OPENAI_API_KEY"),
                temperature=0.1
            )
            Settings.embed_model = OpenAIEmbedding(
                api_key=os.getenv("OPENAI_API_KEY")
            )
            
            # Initialize ChromaDB
            self.chroma_client = chromadb.PersistentClient(
                path="./chroma_db",
                settings=ChromaSettings(anonymized_telemetry=False)
            )
            
            # Create or get collection
            try:
                chroma_collection = self.chroma_client.get_collection("educational_docs")
            except:
                chroma_collection = self.chroma_client.create_collection("educational_docs")
            
            # Create vector store
            vector_store = ChromaVectorStore(chroma_collection=chroma_collection)
            storage_context = StorageContext.from_defaults(vector_store=vector_store)
            
            # Check if we need to load documents
            content_path = Path("./content")
            if content_path.exists() and any(content_path.iterdir()):
                # Load documents
                documents = []
                for file_path in self._get_document_files(content_path):
                    try:
                        doc_text = self._process_document(file_path)
                        if doc_text.strip():
                            # Create document with metadata
                            from llama_index.core import Document
                            doc = Document(
                                text=doc_text,
                                metadata={
                                    "file_path": str(file_path),
                                    "file_name": file_path.name,
                                    "file_type": file_path.suffix,
                                    "category": file_path.parent.name
                                }
                            )
                            documents.append(doc)
                    except Exception as e:
                        logger.error(f"Error loading document {file_path}: {e}")
                
                if documents:
                    # Create index
                    if len(chroma_collection.get()["ids"]) == 0:
                        # New index
                        self.index = VectorStoreIndex.from_documents(
                            documents,
                            storage_context=storage_context
                        )
                        logger.info(f"Created new index with {len(documents)} documents")
                    else:
                        # Existing index
                        self.index = VectorStoreIndex.from_vector_store(
                            vector_store,
                            storage_context=storage_context
                        )
                        logger.info("Loaded existing index")
                else:
                    logger.warning("No valid documents found in content directory")
            else:
                logger.warning("Content directory not found or empty")
            
            # Create query engine if we have an index
            if self.index:
                self.query_engine = self.index.as_query_engine(
                    similarity_top_k=3,
                    response_mode="compact"
                )
            
            logger.info("RAG system initialized successfully")
            
        except Exception as e:
            logger.error(f"Error initializing RAG system: {e}")
        finally:
            self.is_initializing = False
    
    def _get_document_files(self, directory: Path) -> List[Path]:
        """Get all supported document files from directory"""
        supported_extensions = {'.pdf', '.docx', '.txt', '.md'}
        files = []
        
        for item in directory.rglob('*'):
            if item.is_file() and item.suffix.lower() in supported_extensions:
                files.append(item)
        
        return files
    
    def _process_document(self, file_path: Path) -> str:
        """Process a document based on its type"""
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return DocumentProcessor.process_pdf(file_path)
        elif suffix == '.docx':
            return DocumentProcessor.process_docx(file_path)
        elif suffix in ['.txt', '.md']:
            return DocumentProcessor.process_text_file(file_path)
        else:
            logger.warning(f"Unsupported file type: {suffix}")
            return ""
    
    async def _check_rate_limit(self, user_id: int) -> bool:
        """Check if user is within rate limits"""
        try:
            rate_limit_cog = self.bot.get_cog('RateLimitCog')
            if rate_limit_cog:
                return await rate_limit_cog.check_rate_limit(user_id, 'qa_command')
            return True  # If no rate limiting, allow
        except Exception as e:
            logger.error(f"Error checking rate limit: {e}")
            return True
    
    @commands.hybrid_command(name="ask", description="Ask a technical question with RAG-powered answers")
    async def ask(self, ctx, *, question: str):
        """Ask a technical question and get an AI-powered answer with source citations"""
        
        # Check rate limiting
        if not await self._check_rate_limit(ctx.author.id):
            rate_limit_cog = self.bot.get_cog('RateLimitCog')
            if rate_limit_cog:
                retry_after = await rate_limit_cog.get_retry_after(ctx.author.id)
                embed = discord.Embed(
                    title="⏰ Rate Limited",
                    description=f"Please wait {retry_after:.1f} seconds before asking another question.",
                    color=discord.Color.orange()
                )
                await ctx.send(embed=embed, ephemeral=True)
                return
        
        # Check if system is ready
        if not self.index or not self.query_engine:
            embed = discord.Embed(
                title="🔄 System Initializing",
                description="The Q&A system is still initializing. Please try again in a moment.",
                color=discord.Color.yellow()
            )
            await ctx.send(embed=embed)
            return
        
        # Defer the response for longer processing
        await ctx.defer()
        
        try:
            # Get conversation context
            conversation_context = self.conversation_memory.get_context(ctx.author.id)
            
            # Add context to the query if available
            enhanced_query = question
            if conversation_context:
                enhanced_query = f"Previous context:\n{conversation_context}\n\nNew question: {question}"
            
            # Query the RAG system
            response = self.query_engine.query(enhanced_query)
            
            # Extract answer and sources
            answer = str(response)
            sources = []
            
            # Get source information from the response
            if hasattr(response, 'source_nodes'):
                for node in response.source_nodes:
                    if hasattr(node, 'metadata'):
                        file_name = node.metadata.get('file_name', 'Unknown')
                        category = node.metadata.get('category', 'Unknown')
                        sources.append(f"{category}/{file_name}")
            
            # Remove duplicates and limit sources
            sources = list(set(sources))[:3]
            
            # Create response embed
            embed = discord.Embed(
                title="🎓 Educational Assistant",
                description=f"**Question:** {question}",
                color=discord.Color.blue()
            )
            
            # Add answer (split if too long)
            if len(answer) <= 1024:
                embed.add_field(name="📝 Answer", value=answer, inline=False)
            else:
                # Split long answers
                embed.add_field(name="📝 Answer", value=answer[:1021] + "...", inline=False)
                # Send continuation if needed
                if len(answer) > 1021:
                    continuation = answer[1021:]
                    if len(continuation) <= 2000:
                        await ctx.send(f"**Continued:**\n{continuation}")
            
            # Add sources
            if sources:
                sources_text = "\n".join([f"• {source}" for source in sources])
                embed.add_field(name="📚 Sources", value=sources_text, inline=False)
            
            # Add footer
            embed.set_footer(text="💡 Ask follow-up questions for more details!")
            
            # Send response
            await ctx.send(embed=embed)
            
            # Store in conversation memory
            self.conversation_memory.add_interaction(
                ctx.author.id, question, answer, sources
            )
            
            logger.info(f"Q&A query processed for user {ctx.author.id}: {question[:50]}...")
            
        except Exception as e:
            logger.error(f"Error processing Q&A query: {e}")
            
            error_embed = discord.Embed(
                title="❌ Error",
                description="Sorry, I encountered an error while processing your question. Please try again later.",
                color=discord.Color.red()
            )
            
            await ctx.send(embed=error_embed)
    
    @commands.hybrid_command(name="reindex_documents", description="Reindex all documents in the content directory")
    @commands.has_permissions(administrator=True)
    async def reindex_documents(self, ctx):
        """Reindex all documents (Admin only)"""
        await ctx.defer()
        
        try:
            # Clear existing index
            if self.chroma_client:
                try:
                    self.chroma_client.delete_collection("educational_docs")
                except:
                    pass
            
            # Reinitialize the system
            await self.initialize_rag_system()
            
            embed = discord.Embed(
                title="✅ Reindexing Complete",
                description="All documents have been reindexed successfully.",
                color=discord.Color.green()
            )
            await ctx.send(embed=embed)
            
        except Exception as e:
            logger.error(f"Error reindexing documents: {e}")
            embed = discord.Embed(
                title="❌ Reindexing Failed",
                description="An error occurred while reindexing documents.",
                color=discord.Color.red()
            )
            await ctx.send(embed=embed)
    
    @commands.hybrid_command(name="clear_conversation", description="Clear your conversation history")
    async def clear_conversation(self, ctx):
        """Clear conversation memory for the user"""
        self.conversation_memory.clear_conversation(ctx.author.id)
        
        embed = discord.Embed(
            title="🗑️ Conversation Cleared",
            description="Your conversation history has been cleared.",
            color=discord.Color.green()
        )
        await ctx.send(embed=embed, ephemeral=True)
    
    @commands.hybrid_command(name="qa_stats", description="View Q&A system statistics")
    async def qa_stats(self, ctx):
        """Display Q&A system statistics"""
        try:
            embed = discord.Embed(
                title="📊 Q&A System Statistics",
                color=discord.Color.blue()
            )
            
            # System status
            status = "🟢 Online" if self.index else "🔴 Offline"
            embed.add_field(name="System Status", value=status, inline=True)
            
            # Document count
            if self.chroma_client:
                try:
                    collection = self.chroma_client.get_collection("educational_docs")
                    doc_count = len(collection.get()["ids"])
                    embed.add_field(name="Indexed Documents", value=str(doc_count), inline=True)
                except:
                    embed.add_field(name="Indexed Documents", value="Unknown", inline=True)
            
            # Active conversations
            active_conversations = len(self.conversation_memory.conversations)
            embed.add_field(name="Active Conversations", value=str(active_conversations), inline=True)
            
            # Content directory info
            content_path = Path("./content")
            if content_path.exists():
                file_count = len(self._get_document_files(content_path))
                embed.add_field(name="Available Documents", value=str(file_count), inline=True)
            
            await ctx.send(embed=embed)
            
        except Exception as e:
            logger.error(f"Error getting Q&A stats: {e}")
            await ctx.send("Error retrieving statistics.")

async def setup(bot):
    """Setup function to add the cog to the bot"""
    await bot.add_cog(QACog(bot))
    logger.info("QA Cog loaded successfully")
